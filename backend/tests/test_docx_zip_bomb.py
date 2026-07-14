"""Güvenlik denetimi bulgusu (gerçek PoC ile ölçüldü): 628KB'lık hazırlanmış bir
DOCX, 240M karaktere / 540MB+ belleğe genişliyordu çünkü 100.000 karakter metin
sınırı kontrolü extraction TAMAMLANDIKTAN sonra çalışıyordu. Bu dosya hem
düzeltmenin zip bomb'u gerçekten reddettiğini, hem de normal kullanıcıların
gerçek DOCX dosyalarının etkilenmediğini kanıtlar."""

import io
import time
import zipfile

import pytest
from docx import Document

from app.utils.text_extraction import TextExtractionError, extract_text

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_CONTENT_TYPES_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="rels" '
    'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
    '<Default Extension="xml" ContentType="application/xml"/>'
    '<Override PartName="/word/document.xml" '
    'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    "</Types>"
)
_RELS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
    'Target="word/document.xml"/>'
    "</Relationships>"
)
_DOCUMENT_XML_HEADER = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:body><w:p><w:r><w:t>"
)
_DOCUMENT_XML_FOOTER = "</w:t></w:r></w:p></w:body></w:document>"


def _build_zip_bomb_docx(*, repeated_char_count: int) -> bytes:
    """Gerçek bir DOCX zip yapısı, ama word/document.xml'i son derece
    sıkıştırılabilir (tek karakterin milyonlarca tekrarı) dev bir metin bloğu
    içeriyor — yüksek compression-ratio zip bomb'u simüle eder."""
    huge_text = "A" * repeated_char_count
    document_xml = _DOCUMENT_XML_HEADER + huge_text + _DOCUMENT_XML_FOOTER

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)
        archive.writestr("_rels/.rels", _RELS_XML)
        archive.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def _build_legitimate_docx() -> bytes:
    """python-docx ile gerçek, küçük, normal bir CV benzeri belge üretir."""
    document = Document()
    document.add_paragraph("Ayşe Yılmaz")
    document.add_paragraph("Yazılım geliştirme alanında 5 yıllık deneyim.")
    for i in range(20):
        document.add_paragraph(f"Deneyim satırı {i}: proje teslim edildi.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_zip_bomb_docx_is_rejected_before_expanding() -> None:
    """~40MB'a genişleyecek (yaklaşık 400:1 sıkıştırma oranı) bir dosya,
    extraction hiç başlamadan reddedilmeli — hızlı ve düşük bellekle."""
    bomb = _build_zip_bomb_docx(repeated_char_count=40_000_000)
    assert len(bomb) < 200_000, "Test fixture'ı kendisi de küçük/sıkıştırılabilir olmalı"

    start = time.monotonic()
    with pytest.raises(TextExtractionError):
        extract_text(content=bomb, mime_type=_DOCX_MIME)
    elapsed = time.monotonic() - start

    # Gerçek extraction (düzeltme öncesi) 13.5 saniye sürüyordu; reddetme zip
    # merkezi dizinini okumaktan ibaret olduğu için çok daha hızlı olmalı.
    assert elapsed < 2.0, f"Reddetme {elapsed:.2f}s sürdü — hâlâ içerik açılıyor olabilir"


def test_zip_bomb_with_many_small_entries_is_rejected() -> None:
    """Tek büyük dosya yerine çok sayıda entry ile aşırı toplam boyut/entry
    sayısı da reddedilmeli (yalnızca tek-dosya kontrolüne güvenilmemeli)."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)
        archive.writestr("_rels/.rels", _RELS_XML)
        archive.writestr(
            "word/document.xml", _DOCUMENT_XML_HEADER + "A" * 1000 + _DOCUMENT_XML_FOOTER
        )
        for i in range(300):  # varsayılan limit olan 200 entry'yi aşıyor
            archive.writestr(f"word/media/image{i}.bin", b"x" * 100)

    with pytest.raises(TextExtractionError):
        extract_text(content=buffer.getvalue(), mime_type=_DOCX_MIME)


def test_legitimate_small_docx_is_not_affected() -> None:
    """Normal kullanıcıların gerçek CV'leri hiçbir şekilde etkilenmemeli."""
    content = _build_legitimate_docx()
    text = extract_text(content=content, mime_type=_DOCX_MIME)
    assert "Ayşe Yılmaz" in text
    assert "Deneyim satırı 19" in text
